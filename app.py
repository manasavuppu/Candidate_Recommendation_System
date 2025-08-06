import streamlit as st
import numpy as np
import matplotlib.pyplot as plt
from base_model import rank_candidates
from hybrid_model import query_top_k_candidates
import fitz  # PyMuPDF for PDFs
from docx import Document
import io

# ============================
# App Setup
# ============================
st.set_page_config(page_title="Candidate Recommendation Engine", layout="wide")
st.title("Candidate Recommendation Engine")

# ============================
# Sidebar Settings
# ============================
st.sidebar.title("Settings")
model_choice = st.sidebar.radio("Select Matching Model", ["Base Model (SentenceTransformers)", "Hybrid Model"])
use_summaries = st.sidebar.checkbox("Generate AI Fit Summaries", value=False)
top_k = st.sidebar.slider("Number of Top Candidates", min_value=1, max_value=10, value=5)
min_score = float(st.sidebar.slider("Minimum Similarity Score", 0.0, 1.0, 0.5, step=0.01))

# ============================
# Job Description Input
# ============================
st.markdown("### Job Description Input")
job_input_method = st.radio("How would you like to provide the job description?", ["Paste Text", "Upload File"])
job_desc = ""

def extract_text(uploaded_file):
    if uploaded_file.name.endswith(".pdf"):
        doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
        return "\n".join([page.get_text() for page in doc])
    elif uploaded_file.name.endswith(".txt"):
        return uploaded_file.read().decode("utf-8", errors="ignore")
    elif uploaded_file.name.endswith(".docx"):
        doc = Document(uploaded_file)
        return "\n".join([p.text for p in doc.paragraphs])
    else:
        return ""

if job_input_method == "Paste Text":
    job_desc = st.text_area("Paste the job description here", height=200)
else:
    uploaded_jd_file = st.file_uploader("Upload a Job Description File (.txt, .pdf, .docx)", type=["txt", "pdf", "docx"], key="jd_upload")
    if uploaded_jd_file:
        job_desc = extract_text(uploaded_jd_file)
        st.success("Job description loaded from file!")

# ============================
# Resume Upload
# ============================
uploaded_files = st.file_uploader(
    "Upload Candidate Resumes (PDF, DOCX, or TXT)",
    type=["pdf", "txt", "docx"],
    accept_multiple_files=True
)

def extract_text_from_file(uploaded_file):
    try:
        if uploaded_file.name.endswith(".pdf"):
            doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
            return "\n".join([page.get_text() for page in doc])
        elif uploaded_file.name.endswith(".docx"):
            doc = Document(uploaded_file)
            return "\n".join([para.text for para in doc.paragraphs])
        elif uploaded_file.name.endswith(".txt"):
            return uploaded_file.read().decode("utf-8", errors="ignore")
        else:
            return ""
    except Exception as e:
        return f"[ERROR] Could not read {uploaded_file.name}: {str(e)}"

# ============================
# Candidate Ranking
# ============================
if st.button("Find Top Candidates"):
    if not job_desc or not uploaded_files:
        st.warning("Please provide both a job description and at least one resume.")
    else:
        with st.spinner("Analyzing resumes..."):
            resume_data = []
            for file in uploaded_files:
                text = extract_text_from_file(file)
                resume_data.append({"text": text, "file_name": file.name})

            # Use correct model selection match
            if "Base Model" in model_choice:
                results = rank_candidates(
                    job_desc=job_desc,
                    resumes=resume_data,
                    top_k=top_k,
                    include_summary=use_summaries,
                    min_score=min_score
                )
            else:
                results = query_top_k_candidates(
                    job_desc=job_desc,
                    resumes=resume_data,
                    top_k=top_k,
                    include_summary=use_summaries,
                    min_score=min_score
                )

            st.session_state["latest_results"] = results

            # ============================
            # Display Results
            # ============================
            st.subheader(f"Top {top_k} Matches")
            if not results:
                st.warning("No candidates matched the minimum similarity score threshold. Try lowering the threshold.")
            else:
                st.caption(f"{len(results)} out of {len(resume_data)} candidates passed the threshold of {min_score}.")
                for i, res in enumerate(results, 1):
                    name = res.get('file_name', res.get('name', f"Candidate {i}"))
                    score = res.get('similarity_score', res.get('score', 0))
                    st.markdown(f"### {i}. {name}")
                    st.write(f"**Similarity Score:** `{score}`")
                    if use_summaries and "summary" in res:
                        st.info(res['summary'])

# ============================
# Behind-the-Scenes Section
# ============================
if "latest_results" in st.session_state and st.button("Show Behind-the-Scenes (Embeddings & Similarity)"):
    results = st.session_state["latest_results"]
    st.subheader("Resume Content & Embeddings")

    for res in results:
        st.markdown(f"**{res.get('file_name', res.get('name', 'Unnamed'))}** — Similarity: `{res.get('similarity_score')}`")
        if 'embedding' in res:
            emb = np.array(res['embedding'])
            st.text(f"Embedding shape: {emb.shape}")
            st.text(f"First 10 dims: {np.round(emb[:10], 3)}")
        st.markdown("---")

    # Similarity Chart
    st.subheader("Similarity Score Chart")
    fig, ax = plt.subplots(figsize=(10, 4))
    labels = [r.get('file_name', r.get('name', f"Candidate {i+1}")) for i, r in enumerate(results)]
    scores = [r.get('similarity_score', r.get('score', 0)) for r in results]
    ax.bar(labels, scores)
    ax.set_ylabel("Cosine Similarity")
    ax.set_title("Top Candidate Scores")
    plt.xticks(rotation=45, ha='right')
    st.pyplot(fig)

    # Debug Prompts
    if use_summaries:
        st.subheader("Summary Prompt Debug")
        for res in results:
            if 'prompt' in res and 'summary' in res:
                with st.expander(f"{res.get('file_name', res.get('name', 'Unnamed'))} Prompt & Summary"):
                    st.markdown("**Prompt Sent to LLM:**")
                    st.code(res['prompt'])
                    st.markdown("**LLM Output Summary:**")
                    st.write(res['summary'])
